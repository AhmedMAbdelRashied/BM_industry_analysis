import os
import re
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import io
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
 
class DocumentFormatter:
    def __init__(self, filename, selected_language):
        """Initialize the DocumentFormatter class with a given file name and language."""
        self.filename = filename
        self.selected_language = selected_language  # Store the selected language
        if os.path.exists(self.filename):
            os.remove(self.filename)
        self.doc = Document()
        self.in_visualization_block = False
        self.visualization_code = []
        self.first_title_added = False  # Flag to track first title occurrence
 
    def set_paragraph_alignment(self, paragraph):
        """Sets the paragraph alignment based on the selected language."""
        paragraph_format = paragraph.paragraph_format
        if self.selected_language == 'Arabic':
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p = paragraph._element
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            p.get_or_add_pPr().append(bidi)
        else:
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
 
    def _set_font_to_dubai(self, run):
        """Sets the font of the given run to Dubai."""
        run.font.name = 'Dubai'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Dubai')  # Ensure it's applied for all regions
 
    def add_text_to_doc(self, text):
        """Adds text to the document, identifying and formatting headers, tables, and bold text."""
        lines = text.split('\n')
        table = None
        in_table = False
        current_table_columns = 0  # Track the number of columns in the current table
 
        # Updated regex pattern to handle both **bold** and *bold* text
        bold_pattern = re.compile(r'(\*\*(.*?)\*\*)|(\*(.*?)\*)')
 
        for line in lines:
            # Start Of visualization mechanism
            # Handle ```python
            if '```python' in line:
                line = line.replace('```python', '')
            if '```markdown' in line:
                line = line.replace('```markdown', '')
            if '```' in line:
                line = line.replace('```', '')
 
            # Handle visualization tags
            if '[visualization]' in line:
                self.in_visualization_block = True
                self.visualization_code = []
                continue
            elif '[/visualization]' in line:
                self.in_visualization_block = False
                self._execute_and_insert_visualization()
                continue
 
            # Collect code for visualization if inside a block
            if self.in_visualization_block:
                self.visualization_code.append(line)
                continue
            # End Of visualization mechanism
 
            # Check if the line is a table separator and skip it
            if re.match(r'^\|\s*[-:]+\s*\|', line):
                continue
 
            # Check if the line is a header like (#)
            if re.match(r'^#\s+', line):
                if not self.first_title_added:
                    # Add the first title and mark the flag
                    line = line.replace('#', '').strip()
                    heading = self.doc.add_heading(line, 0)
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in heading.runs:
                        self._set_font_to_dubai(run)
                    self.first_title_added = True
                else:
                    # Convert further `#` to level 2 headings
                    self._add_header(line, 'Heading 2', 16)
            elif re.match(r'^##\s+', line):
                self._add_header(line, 'Heading 2', 16)
            elif re.match(r'^###\s+', line):
                self._add_header(line, 'Heading 3', 14)
            elif re.match(r'^####\s+', line):
                self._add_header(line, 'Heading 4', 12)
            elif re.match(r'^#####\s+', line):
                self._add_header(line, 'Heading 5', 10)
            elif re.match(r'^######\s+', line):
                self._add_header(line, 'Heading 5', 10)
            # Check if the line is part of a table
            elif '|' in line and line.strip().startswith('|'):
                # Parse the table row
                row_data = [cell.strip() or '-' for cell in line.strip().strip('|').split('|')]
                num_columns = len(row_data)
 
                if not in_table:
                    # Start a new table
                    table = self.doc.add_table(rows=1, cols=num_columns)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for idx, cell_text in enumerate(row_data):
                        self._add_run_with_bold(hdr_cells[idx].paragraphs[0], cell_text, bold_pattern)
                        self.set_paragraph_alignment(hdr_cells[idx].paragraphs[0])
                    current_table_columns = num_columns
                    in_table = True
                else:
                    if num_columns == current_table_columns:
                        # Add a row to the existing table
                        row_cells = table.add_row().cells
                        for idx, cell_text in enumerate(row_data):
                            self._add_run_with_bold(row_cells[idx].paragraphs[0], cell_text, bold_pattern)
                            self.set_paragraph_alignment(row_cells[idx].paragraphs[0])
                    else:
                        # Different number of columns, start a new table
                        # Reset table-related variables
                        table = None
                        in_table = False
                        current_table_columns = 0
 
                        # Start the new table with the current row
                        table = self.doc.add_table(rows=1, cols=num_columns)
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        for idx, cell_text in enumerate(row_data):
                            self._add_run_with_bold(hdr_cells[idx].paragraphs[0], cell_text, bold_pattern)
                            self.set_paragraph_alignment(hdr_cells[idx].paragraphs[0])
                        current_table_columns = num_columns
                        in_table = True
            else:
                if in_table:
                    # End the current table if a non-table line is encountered
                    in_table = False
                    table = None
                    current_table_columns = 0
                if line.strip():
                    # Avoid adding empty paragraphs
                    self._add_paragraph_with_bold(line, bold_pattern)
 
        # After processing all lines, ensure the table is properly closed
        if in_table:
            in_table = False
            table = None
            current_table_columns = 0
 
    def _add_header(self, line, heading_style, font_size):
        """Adds a header to the document with the appropriate style and font size."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(line.strip('# ').strip())
        run.bold = True
        run.font.size = Pt(font_size)
        self._set_font_to_dubai(run)
        paragraph.style = heading_style
        self.set_paragraph_alignment(paragraph)
 
    def _add_paragraph_with_bold(self, line, bold_pattern):
        """Adds a paragraph to the document, handling bold formatting."""
        paragraph = self.doc.add_paragraph()
        self.set_paragraph_alignment(paragraph)
        self._add_run_with_bold(paragraph, line, bold_pattern)
 
    def _add_run_with_bold(self, paragraph, text, bold_pattern):
        """Handles adding text with bold formatting to paragraphs or table cells."""
        last_pos = 0
        for match in bold_pattern.finditer(text):
            if match.group(1):  # **bold**
                start, end = match.span(1)
                if start > last_pos:
                    regular_run = paragraph.add_run(text[last_pos:start])
                    self._set_font_to_dubai(regular_run)
                bold_run = paragraph.add_run(match.group(2))
                bold_run.bold = True
                self._set_font_to_dubai(bold_run)
                last_pos = end
            elif match.group(3):  # *bold*
                start, end = match.span(3)
                if start > last_pos:
                    regular_run = paragraph.add_run(text[last_pos:start])
                    self._set_font_to_dubai(regular_run)
                bold_run = paragraph.add_run(match.group(4))
                bold_run.bold = True
                self._set_font_to_dubai(bold_run)
                last_pos = end
        if last_pos < len(text):
            regular_run = paragraph.add_run(text[last_pos:])
            self._set_font_to_dubai(regular_run)
 
    def _execute_and_insert_visualization(self):
        """Executes the Python code stored for visualization and inserts the resulting graph."""
        code = "\n".join(self.visualization_code)
        try:
            # Create a local namespace for the code
            local_namespace = {}
            # Execute the code
            exec(code, {}, local_namespace)
            # If a plot exists, save it to a stream and insert it into the document
            if plt.get_fignums():
                image_stream = io.BytesIO()
                plt.savefig(image_stream, format='png')
                plt.close()  # Close the plot to free memory
                image_stream.seek(0)
                self.doc.add_picture(image_stream, width=Inches(5))
                image_stream.close()
        except Exception as e:
            print(f"Error executing visualization code: {e}")
 
    def save(self):
        """Saves the document to a file."""
        print('\n---------------------\nDocument Saved.\n---------------------\n')
        self.doc.save(f'{self.filename}')

if __name__ == "__main__":
    sample_text = """# Document Title
    This is an introductory paragraph with some **bold text**.
        | Header1 | Header2 |
        |---------|---------|
        | Cell1   | Cell2   |
        | Cell3   | Cell4   |
        | Header A | Header B | Header C |
        |----------|----------|----------|
        | Data1    | Data2    | Data3    |
        | Data4    | Data5    | Data6    |
        | Data7    | Data8    | Data9    |
        
        Final concluding paragraph.
        """
    
    formatter = DocumentFormatter("output.docx", selected_language="English")
    formatter.add_text_to_doc(sample_text)
    formatter.save()